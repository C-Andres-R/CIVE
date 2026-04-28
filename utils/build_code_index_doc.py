"""Genera un índice navegable por módulo e historia de usuario del proyecto."""

from __future__ import annotations

import ast
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Indice_Historias_Usuario_CIVE.docx"


@dataclass
class CodeRow:
    """Representa una función, método o clase con su ubicación."""

    file_path: Path
    kind: str
    name: str
    description: str
    start_line: int
    end_line: int


@dataclass
class UserStory:
    """Representa una historia de usuario con archivos relacionados."""

    story_id: str
    title: str
    files: list[str]


@dataclass
class ModuleGroup:
    """Representa un módulo funcional del sistema."""

    module_id: str
    name: str
    description: str
    stories: list[UserStory]


MODULES: list[ModuleGroup] = [
    ModuleGroup(
        module_id="Mod-00",
        name="Administración y Control de Usuarios",
        description=(
            "El sistema contará con un módulo para gestionar usuarios, roles, permisos, estados "
            "de cuenta (activación/desactivación) y mecanismos de autenticación segura con roles "
            "específicos claramente definidos."
        ),
        stories=[
            UserStory("HU-000", "Gestionar Roles y Permisos", [
                "app/routes/usuarios.py",
                "app/models/rol.py",
                "app/models/permiso.py",
                "app/models/rol_permiso.py",
                "app/auth/decorators.py",
            ]),
            UserStory("HU-000", "Gestionar Usuarios", [
                "app/routes/usuarios.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-000", "Autenticar de forma segura", [
                "app/routes/pages.py",
                "app/auth/service.py",
                "app/auth/password_policy.py",
                "app/security.py",
                "app/auth/decorators.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 001",
        name="Chatbot Inteligente",
        description="Chatbot para resolver dudas, agendar citas y enviar recordatorios.",
        stories=[
            UserStory("HU-001", "Consultar de forma Automática (FAQs)", [
                "app/routes/chat.py",
                "app/models/chatbot_faq.py",
            ]),
            UserStory("HU-002", "Solicitar Citas desde Chatbot", [
                "app/routes/chat.py",
                "app/routes/citas.py",
                "app/models/cita.py",
            ]),
            UserStory("HU-003", "Generar recordatorios automáticos mediante Chatbot", [
                "app/routes/chat.py",
                "app/routes/citas.py",
                "app/models/recordatorio_cita.py",
            ]),
            UserStory("HU-004", "Evaluar el servicio mediante chatbot", [
                "app/routes/chat.py",
                "app/routes/encuestas.py",
                "app/models/encuesta_satisfaccion.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 002",
        name="Gestión de Citas",
        description=(
            "Permite crear, modificar, cancelar y visualizar citas, con historial completo "
            "y recordatorios automáticos por correo electrónico."
        ),
        stories=[
            UserStory("HU-005", "Crear citas", [
                "app/routes/citas.py",
                "app/models/cita.py",
            ]),
            UserStory("HU-006", "Modificar y cancelar de citas", [
                "app/routes/citas.py",
                "app/models/cita.py",
            ]),
            UserStory("HU-007", "Visualizar y filtrar citas de forma avanzada", [
                "app/routes/citas.py",
            ]),
            UserStory("HU-008", "Generar recordatorios automáticos", [
                "app/routes/citas.py",
                "app/models/recordatorio_cita.py",
            ]),
            UserStory("HU-009", "Consultar disponibilidad de veterinarios", [
                "app/routes/citas.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-010", "Reasignar citas por ausencia de forma automática", [
                "app/routes/citas.py",
                "app/models/cita.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 003",
        name="Módulo de registro de mascotas",
        description=(
            "Permite registrar, actualizar, inactivar mascotas y mantener el historial completo "
            "vinculado al dueño."
        ),
        stories=[
            UserStory("HU-011", "Registrar mascota nueva", [
                "app/routes/mascotas.py",
                "app/models/mascota.py",
            ]),
            UserStory("HU-012", "Actualizar datos mascota", [
                "app/routes/mascotas.py",
                "app/models/mascota.py",
            ]),
            UserStory("HU-013", "Inactivar mascota", [
                "app/routes/mascotas.py",
                "app/models/mascota.py",
            ]),
            UserStory("HU-014", "Visualizar historial completo mascota", [
                "app/routes/mascotas.py",
                "app/routes/expedientes.py",
                "app/models/mascota.py",
            ]),
            UserStory("HU-015", "Asociar mascota-dueño", [
                "app/routes/mascotas.py",
                "app/routes/clientes.py",
                "app/models/mascota.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-016", "Mostrar galería multimedia mascota", [
                "app/routes/mascotas.py",
                "app/models/foto_mascota.py",
                "app/models/documento_mascota.py",
            ]),
            UserStory("HU-017", "Registrar comportamientos especiales", [
                "app/routes/mascotas.py",
                "app/models/mascota.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 004",
        name="Gestión de Dueños (Clientes)",
        description=(
            "El módulo permite crear, actualizar, inactivar clientes y mantener la información "
            "clara, con asociación directa a sus mascotas."
        ),
        stories=[
            UserStory("HU-018", "Registrar clientes", [
                "app/routes/clientes.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-019", "Actualizar datos del cliente", [
                "app/routes/clientes.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-020", "Inactivar cliente", [
                "app/routes/clientes.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-021", "Generar notificaciones automatizadas", [
                "app/routes/clientes.py",
                "app/routes/citas.py",
                "app/models/recordatorio_cita.py",
            ]),
            UserStory("HU-022", "Visualizar mascotas asociadas", [
                "app/routes/clientes.py",
                "app/routes/mascotas.py",
                "app/models/mascota.py",
            ]),
            UserStory("HU-023", "Generar historial financiero del cliente", [
                "app/routes/clientes.py",
                "app/models/facturacion.py",
            ]),
            UserStory("HU-024", "Mostrar portal del cliente", [
                "app/routes/clientes.py",
                "app/routes/pages.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 005",
        name="Gestión de Expedientes Médicos y Tratamientos",
        description=(
            "El módulo gestiona el historial clínico completo por mascota, incluyendo consultas, "
            "tratamientos, vacunas y alergias."
        ),
        stories=[
            UserStory("HU-025", "Registrar consulta médica", [
                "app/routes/expedientes.py",
                "app/models/consulta_medica.py",
            ]),
            UserStory("HU-026", "Actualizar expediente médico", [
                "app/routes/expedientes.py",
                "app/models/consulta_medica.py",
            ]),
            UserStory("HU-027", "Visualizar historial médico completo", [
                "app/routes/expedientes.py",
                "app/models/consulta_medica.py",
                "app/models/mascota.py",
            ]),
            UserStory("HU-028", "Descargar e imprimir reportes clínicos", [
                "app/routes/expedientes.py",
            ]),
            UserStory("HU-029", "Registrar vacunas y alergias", [
                "app/routes/expedientes.py",
                "app/models/vacuna_alergia.py",
            ]),
            UserStory("HU-030", "Generar notificaciones automáticas por tratamiento pendiente", [
                "app/followups.py",
                "app/routes/expedientes.py",
                "app/models/seguimiento_tratamiento.py",
            ]),
            UserStory("HU-031", "Registrar análisis clínicos detallados", [
                "app/routes/expedientes.py",
                "app/models/analisis_clinico.py",
            ]),
            UserStory("HU-032", "Controlar inventario de medicamentos", [
                "app/routes/expedientes.py",
                "app/models/insumo_clinico.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 006",
        name="Generación de reportes",
        description=(
            "El módulo permite generar reportes administrativos, financieros y de productividad "
            "en múltiples formatos descargables."
        ),
        stories=[
            UserStory("HU-033", "Generar Reporte Administrativo de Citas", [
                "app/routes/reportes.py",
                "app/models/cita.py",
            ]),
            UserStory("HU-034", "Generar Reporte Financiero Ingresos", [
                "app/routes/reportes.py",
                "app/models/facturacion.py",
            ]),
            UserStory("HU-035", "Generar Reporte Productividad Veterinarios", [
                "app/routes/reportes.py",
                "app/models/cita.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-036", "Generar Reportes en formatos descargables", [
                "app/routes/reportes.py",
            ]),
            UserStory("HU-037", "Generar Reporte mensual de clientes nuevos", [
                "app/routes/reportes.py",
                "app/models/usuario.py",
            ]),
            UserStory("HU-038", "Generar Reporte de medicamentos más utilizados", [
                "app/routes/reportes.py",
                "app/models/insumo_clinico.py",
            ]),
            UserStory("HU-039", "Generar Reporte clientes con mayor frecuencia de visitas", [
                "app/routes/reportes.py",
                "app/models/cita.py",
                "app/models/usuario.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 007",
        name="Análisis y Visualización de Datos",
        description=(
            "Este módulo permitirá analizar estadísticamente datos relevantes del sistema y "
            "visualizar tendencias mediante gráficos dinámicos."
        ),
        stories=[
            UserStory("HU-040", "Analizar frecuencia y tipo de consultas", [
                "app/routes/datos.py",
                "app/models/cita.py",
                "app/models/consulta_medica.py",
            ]),
            UserStory("HU-041", "Visualizar Tendencias en servicios solicitados", [
                "app/routes/datos.py",
            ]),
            UserStory("HU-042", "Generar Visualizaciones gráficas dinámicas", [
                "app/routes/datos.py",
            ]),
            UserStory("HU-043", "Mostrar Panel administrativo con estadísticas rápidas", [
                "app/routes/datos.py",
                "app/routes/pages.py",
            ]),
            UserStory("HU-044", "Realizar Análisis predictivo de citas futuras", [
                "app/routes/datos.py",
                "app/models/cita.py",
            ]),
            UserStory("HU-045", "Realizar Monitoreo en tiempo real de indicadores clave", [
                "app/routes/datos.py",
            ]),
        ],
    ),
    ModuleGroup(
        module_id="Mod 008",
        name="Encuestas de Satisfacción y Retroalimentación",
        description=(
            "Permite automatizar el envío de encuestas de satisfacción tras consultas o "
            "tratamientos, recopilando y analizando los resultados."
        ),
        stories=[
            UserStory("HU-046", "Automatizar envío encuestas", [
                "app/routes/encuestas.py",
                "app/models/encuesta_satisfaccion.py",
            ]),
            UserStory("HU-047", "Responder encuestas", [
                "app/routes/encuestas.py",
                "app/models/encuesta_satisfaccion.py",
                "app/models/encuesta_pregunta.py",
            ]),
            UserStory("HU-048", "Generar reportes de satisfacción", [
                "app/routes/encuestas.py",
                "app/models/encuesta_satisfaccion.py",
            ]),
            UserStory("HU-049", "Visualizar resultados de forma gráfica", [
                "app/routes/encuestas.py",
            ]),
            UserStory("HU-050", "Analizar comentarios recibidos", [
                "app/routes/encuestas.py",
                "app/models/encuesta_satisfaccion.py",
            ]),
        ],
    ),
]


def _humanize(name: str) -> str:
    """Convierte nombres técnicos a una frase simple en español."""
    cleaned = name.strip("_") or "bloque"
    cleaned = re.sub(r"(?<!^)([A-Z])", r" \1", cleaned)
    words = [part.lower() for part in re.split(r"[_\s]+", cleaned) if part]
    return " ".join(words)


def _summary_from_name(kind: str, name: str) -> str:
    """Genera una descripción breve cuando no hay docstring clara."""
    phrase = _humanize(name)
    replacements = [
        ("validar ", "Valida "),
        ("obtener ", "Obtiene "),
        ("guardar ", "Guarda "),
        ("eliminar ", "Elimina "),
        ("parsear ", "Convierte "),
        ("construir ", "Construye "),
        ("build ", "Construye "),
        ("sincronizar ", "Sincroniza "),
        ("redirigir ", "Redirige "),
        ("generar ", "Genera "),
        ("enviar ", "Envía "),
        ("formatear ", "Da formato a "),
        ("render ", "Renderiza "),
        ("require ", "Valida "),
        ("mark ", "Marca "),
        ("clear ", "Limpia "),
        ("touch ", "Actualiza "),
        ("is ", "Verifica si "),
    ]
    for source, target in replacements:
        if phrase.startswith(source):
            return f"{target}{phrase[len(source):]}."
    prefix = "Define la clase" if kind == "class" else "Función para"
    return f"{prefix} {phrase}."


def _summary_from_docstring(kind: str, name: str, docstring: str | None) -> str:
    """Obtiene una descripción breve desde el docstring."""
    if docstring:
        first_line = next((line.strip() for line in docstring.splitlines() if line.strip()), "")
        if first_line:
            return first_line
    return _summary_from_name(kind, name)


def _iter_class_methods(file_path: Path, node: ast.ClassDef) -> list[CodeRow]:
    """Extrae métodos declarados dentro de una clase."""
    methods: list[CodeRow] = []
    for child in node.body:
        if isinstance(child, (ast.FunctionDef, ast.AsyncFunctionDef)):
            methods.append(
                CodeRow(
                    file_path=file_path,
                    kind="method",
                    name=f"{node.name}.{child.name}",
                    description=_summary_from_docstring("function", child.name, ast.get_docstring(child)),
                    start_line=child.lineno,
                    end_line=getattr(child, "end_lineno", child.lineno),
                )
            )
    return methods


def _iter_code_items(file_path: Path, tree: ast.AST) -> list[CodeRow]:
    """Extrae clases, métodos y funciones del archivo."""
    rows: list[CodeRow] = []
    for node in tree.body:
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
            rows.append(
                CodeRow(
                    file_path=file_path,
                    kind="function",
                    name=node.name,
                    description=_summary_from_docstring("function", node.name, ast.get_docstring(node)),
                    start_line=node.lineno,
                    end_line=getattr(node, "end_lineno", node.lineno),
                )
            )
        elif isinstance(node, ast.ClassDef):
            rows.append(
                CodeRow(
                    file_path=file_path,
                    kind="class",
                    name=node.name,
                    description=_summary_from_docstring("class", node.name, ast.get_docstring(node)),
                    start_line=node.lineno,
                    end_line=getattr(node, "end_lineno", node.lineno),
                )
            )
            rows.extend(_iter_class_methods(file_path, node))
    return rows


def collect_code_index() -> dict[str, list[CodeRow]]:
    """Recorre el proyecto y construye un índice por archivo."""
    files = [Path("run.py"), Path("wsgi.py")]
    files.extend(sorted(Path("app").rglob("*.py")))
    files.extend(sorted(Path("tests").rglob("*.py")))

    collected: dict[str, list[CodeRow]] = {}
    for rel_path in files:
        full_path = ROOT / rel_path
        tree = ast.parse(full_path.read_text(encoding="utf-8"))
        collected[str(rel_path)] = _iter_code_items(rel_path, tree)
    return collected


def _bookmark_paragraph(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    """Ancla un párrafo para navegación interna."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_link(paragraph, text: str, anchor: str) -> None:
    """Inserta un enlace interno hacia un bookmark."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "Hyperlink")
    rpr.append(rstyle)
    run.append(rpr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Escribe texto en una celda con formato uniforme."""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold


def _story_anchor(module: ModuleGroup, story: UserStory) -> str:
    """Genera un identificador interno único por historia."""
    slug = re.sub(r"[^a-z0-9]+", "-", story.title.lower()).strip("-")
    return f"{module.module_id.lower().replace(' ', '').replace('-', '')}_{story.story_id.lower()}_{slug}"[:40]


def _rows_for_story(story: UserStory, code_index: dict[str, list[CodeRow]]) -> list[CodeRow]:
    """Obtiene las funciones y clases asociadas a una historia."""
    rows: list[CodeRow] = []
    for file_name in story.files:
        rows.extend(code_index.get(file_name, []))
    rows.sort(key=lambda item: (str(item.file_path), item.start_line, item.name))
    return rows


def build_doc(code_index: dict[str, list[CodeRow]]) -> Path:
    """Construye el documento Word por módulo e historia de usuario."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Índice de Historias de Usuario del Proyecto CIVE")
    title_run.bold = True
    title_run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Documento de consulta rápida para ubicar módulos, historias de usuario y funciones relacionadas durante la defensa."
    )

    note = doc.add_paragraph()
    note.add_run("Nota: ").bold = True
    note.add_run(
        "Cada historia de usuario enlaza a una tabla con los archivos más relacionados y el rango exacto de líneas donde se implementa la lógica."
    )

    doc.add_paragraph("")
    index_heading = doc.add_paragraph()
    index_heading.add_run("Índice por módulo e historia de usuario").bold = True
    _bookmark_paragraph(index_heading, "indice_principal", 1)

    bookmarks: dict[str, str] = {}
    bookmark_id = 2
    for module in MODULES:
        module_paragraph = doc.add_paragraph()
        module_paragraph.add_run(f"{module.module_id}: {module.name}").bold = True
        for story in module.stories:
            anchor = _story_anchor(module, story)
            bookmarks[f"{module.module_id}|{story.story_id}|{story.title}"] = anchor
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            _add_internal_link(p, f"{story.story_id}: {story.title}", anchor)
            bookmark_id += 1

    doc.add_page_break()

    for module in MODULES:
        module_heading = doc.add_paragraph()
        module_heading.style = doc.styles["Heading 1"]
        module_heading.add_run(f"{module.module_id}: {module.name}")

        module_desc = doc.add_paragraph()
        module_desc.add_run("Descripción del módulo: ").bold = True
        module_desc.add_run(module.description)

        for story in module.stories:
            anchor_key = f"{module.module_id}|{story.story_id}|{story.title}"
            story_heading = doc.add_paragraph()
            story_heading.style = doc.styles["Heading 2"]
            story_heading.add_run(f"{story.story_id}: {story.title}")
            _bookmark_paragraph(story_heading, bookmarks[anchor_key], bookmark_id)
            bookmark_id += 1

            files_paragraph = doc.add_paragraph()
            files_paragraph.add_run("Archivos relacionados: ").bold = True
            files_paragraph.add_run(", ".join(story.files))

            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            headers = ["Archivo", "Función / clase", "Descripción", "Ubicación específica"]
            for cell, header in zip(table.rows[0].cells, headers):
                _set_cell_text(cell, header, bold=True)

            rows = _rows_for_story(story, code_index)
            if rows:
                for item in rows:
                    row = table.add_row().cells
                    if item.kind == "class":
                        kind_label = "clase"
                    elif item.kind == "method":
                        kind_label = "método"
                    else:
                        kind_label = "función"
                    _set_cell_text(row[0], str(item.file_path))
                    _set_cell_text(row[1], f"{item.name} ({kind_label})")
                    _set_cell_text(row[2], item.description)
                    _set_cell_text(row[3], f"Líneas {item.start_line} a {item.end_line}")
            else:
                row = table.add_row().cells
                _set_cell_text(row[0], "-")
                _set_cell_text(row[1], "Sin coincidencias")
                _set_cell_text(row[2], "No se encontraron funciones relacionadas con esta historia.")
                _set_cell_text(row[3], "-")

            back = doc.add_paragraph()
            _add_internal_link(back, "Volver al índice", "indice_principal")
            doc.add_paragraph("")

    doc.save(OUTPUT)
    return OUTPUT


def main() -> None:
    """Genera el documento final en la raíz del proyecto."""
    code_index = collect_code_index()
    output = build_doc(code_index)
    print(output)


if __name__ == "__main__":
    main()
